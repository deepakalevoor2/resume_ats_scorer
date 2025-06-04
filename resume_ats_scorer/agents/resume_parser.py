import logging
from typing import Dict, Any, Optional
from pathlib import Path
import PyPDF2
import docx
from bs4 import BeautifulSoup
from crewai import Agent
from ..models.schemas import ParsedResume, ResumeSection, FileType
from ..core.exceptions import (
    FileValidationError,
    ParsingError,
    UnsupportedFileTypeError,
    FileNotFoundError
)

logger = logging.getLogger(__name__)


class ResumeParser:
    """Agent responsible for parsing different resume file formats and extracting structured content."""
    
    SUPPORTED_FILE_TYPES = {
        FileType.PDF: ['.pdf'],
        FileType.DOCX: ['.docx', '.doc'],
        FileType.HTML: ['.html', '.htm'],
        FileType.TXT: ['.txt']
    }
    
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    def __init__(self):
        self.agent = Agent(
            role="Resume Parser",
            goal="Extract structured content from resume files",
            backstory="I am an expert at parsing resume files in various formats and extracting their content.",
            verbose=True,
            allow_delegation=False
        )
        
    async def parse_resume(self, file_path: str, file_type: Optional[FileType] = None) -> ParsedResume:
        """Parse a resume file and extract structured content."""
        logger.info(f"Starting resume parsing for file: {file_path}")
        
        try:
            # Validate file
            self._validate_file(file_path, file_type)
            
            # Extract text
            raw_text = self._extract_text(file_path, file_type or self._detect_file_type(file_path))
            if not raw_text.strip():
                raise ParsingError("No text content found in the resume")
                
            # Process content
            sections = self._identify_sections(raw_text)
            keywords = self._extract_keywords(raw_text)
            
            logger.info(f"Successfully parsed resume with {len(sections)} sections and {len(keywords)} keywords")
            
            return ParsedResume(
                raw_text=raw_text,
                sections=sections,
                keywords=keywords,
                metadata={
                    "file_type": file_type,
                    "file_path": file_path,
                    "sections_found": list(sections.keys()),
                    "keyword_count": len(keywords)
                }
            )
            
        except FileNotFoundError as e:
            logger.error(f"File not found: {str(e)}")
            raise
        except FileValidationError as e:
            logger.error(f"File validation failed: {str(e)}")
            raise
        except ParsingError as e:
            logger.error(f"Parsing error: {str(e)}")
            raise
        except Exception as e:
            logger.error(f"Unexpected error during resume parsing: {str(e)}")
            raise ParsingError(f"Failed to parse resume: {str(e)}")
    
    def _validate_file(self, file_path: str, file_type: Optional[FileType] = None) -> None:
        """Validate the resume file before processing."""
        if not Path(file_path).exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        file_size = Path(file_path).stat().st_size
        if file_size > self.MAX_FILE_SIZE:
            raise FileValidationError(f"File size {file_size} bytes exceeds maximum allowed size of {self.MAX_FILE_SIZE} bytes")
            
        if file_type:
            if file_type not in self.SUPPORTED_FILE_TYPES:
                raise UnsupportedFileTypeError(f"Unsupported file type: {file_type}")
            if not any(Path(file_path).suffix.lower() in extensions for extensions in self.SUPPORTED_FILE_TYPES.values()):
                raise FileValidationError(f"File extension does not match specified type: {file_type}")
    
    def _detect_file_type(self, file_path: str) -> FileType:
        """Detect file type based on extension."""
        extension = Path(file_path).suffix.lower()
        for file_type, extensions in self.SUPPORTED_FILE_TYPES.items():
            if extension in extensions:
                return file_type
        raise UnsupportedFileTypeError(f"Unsupported file extension: {extension}")
    
    def _extract_text(self, file_path: str, file_type: FileType) -> str:
        """Extract raw text from resume file based on file type."""
        try:
            if file_type == FileType.PDF:
                return self._extract_from_pdf(file_path)
            elif file_type == FileType.DOCX:
                return self._extract_from_docx(file_path)
            elif file_type == FileType.HTML:
                return self._extract_from_html(file_path)
            elif file_type == FileType.TXT:
                return self._extract_from_txt(file_path)
            else:
                raise UnsupportedFileTypeError(f"Unsupported file type: {file_type}")
        except Exception as e:
            raise ParsingError(f"Error extracting text from {file_type} file: {str(e)}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                if len(reader.pages) == 0:
                    raise ParsingError("PDF file contains no pages")
                for page in reader.pages:
                    page_text = page.extract_text()
                    if not page_text.strip():
                        logger.warning("Empty page detected in PDF")
                    text += page_text
            return text
        except PyPDF2.PdfReadError as e:
            raise ParsingError(f"Invalid PDF file: {str(e)}")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(file_path)
            if not doc.paragraphs:
                raise ParsingError("DOCX file contains no text")
            return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        except Exception as e:
            raise ParsingError(f"Error reading DOCX file: {str(e)}")
    
    def _extract_from_html(self, file_path: str) -> str:
        """Extract text from HTML file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                text = soup.get_text(separator="\n")
                if not text.strip():
                    raise ParsingError("HTML file contains no text content")
                return text
        except Exception as e:
            raise ParsingError(f"Error parsing HTML file: {str(e)}")
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                if not text.strip():
                    raise ParsingError("TXT file is empty")
                return text
        except UnicodeDecodeError:
            raise ParsingError("TXT file contains invalid characters")
        except Exception as e:
            raise ParsingError(f"Error reading TXT file: {str(e)}")
    
    def _identify_sections(self, text: str) -> Dict[ResumeSection, str]:
        """Identify and extract different sections from resume text."""
        sections = {}
        
        # Define section patterns with multiple possible headers
        section_patterns = {
            ResumeSection.EXPERIENCE: ["EXPERIENCE", "WORK EXPERIENCE", "PROFESSIONAL EXPERIENCE"],
            ResumeSection.EDUCATION: ["EDUCATION", "ACADEMIC BACKGROUND"],
            ResumeSection.SKILLS: ["SKILLS", "TECHNICAL SKILLS", "CORE COMPETENCIES"],
            ResumeSection.SUMMARY: ["SUMMARY", "PROFILE", "PROFESSIONAL SUMMARY"]
        }
        
        for section, headers in section_patterns.items():
            content = self._extract_section(text, headers)
            if content:
                sections[section] = content
                logger.debug(f"Found section: {section.value}")
        
        if not sections:
            logger.warning("No standard sections found in resume")
            
        return sections
    
    def _extract_section(self, text: str, section_headers: list) -> str:
        """Extract a specific section's content based on section headers."""
        text_upper = text.upper()
        
        for header in section_headers:
            header_upper = header.upper()
            if header_upper in text_upper:
                start_idx = text_upper.find(header_upper)
                next_section_idx = float('inf')
                
                # Find the next section header
                for section in ResumeSection:
                    section_upper = section.value.upper()
                    if section_upper != header_upper:
                        next_idx = text_upper.find(section_upper, start_idx + len(header))
                        if next_idx != -1 and next_idx < next_section_idx:
                            next_section_idx = next_idx
                
                if next_section_idx == float('inf'):
                    return text[start_idx:].strip()
                else:
                    return text[start_idx:next_section_idx].strip()
        
        return ""
    
    def _extract_keywords(self, text: str) -> list:
        """Extract key skills and keywords from resume text."""
        try:
            # Basic text cleaning
            text = text.lower()
            words = text.replace(',', ' ').replace(';', ' ').split()
            
            # Filter and process keywords
            keywords = [
                word.strip() for word in words 
                if len(word) > 3 and not word.isdigit()
            ]
            
            # Remove duplicates while preserving order
            seen = set()
            unique_keywords = [x for x in keywords if not (x in seen or seen.add(x))]
            
            logger.debug(f"Extracted {len(unique_keywords)} unique keywords")
            return unique_keywords
            
        except Exception as e:
            logger.error(f"Error extracting keywords: {str(e)}")
            return []
